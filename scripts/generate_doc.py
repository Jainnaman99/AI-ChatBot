"""
Generate comprehensive project documentation Word document.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    h.paragraph_format.space_after = Pt(4)
    return h


def add_para(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_code_block(doc, text):
    """Add a monospace code/diagram block."""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), '4')
        bdr.set(qn('w:space'), '4')
        bdr.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bdr)
    pPr.append(pBdr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    p.paragraph_format.left_indent = Inches(0.25 + 0.2 * level)
    p.paragraph_format.space_after = Pt(2)
    return p


def build_doc():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover page ────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('Ministry of Culture India\nRAG Chatbot — Full Technical Architecture')
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x1a, 0x37, 0x6c)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(
        'End-to-End Documentation: Data Ingestion, Vector RAG, LLMs, APIs\n'
        'For Technical Review / Interview Preparation'
    )
    sub_run.italic = True
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # ── Table of Contents placeholder ─────────────────────────────────────────
    add_heading(doc, 'Table of Contents', level=1)
    toc_items = [
        '1. Project Overview',
        '2. System Architecture — High-Level Flow',
        '3. Data Sources',
        '4. Ingestion Pipeline (Crawl → Chunk → Embed → Store)',
        '   4.1 Web Scraping',
        '   4.2 Text Processing & Chunking',
        '   4.3 Embedding Generation',
        '   4.4 Vector Storage (ChromaDB)',
        '5. RAG Query-Time Pipeline',
        '   5.1 Query Entry Points (API)',
        '   5.2 Language Detection & Translation',
        '   5.3 Vector Search & Similarity Scoring',
        '   5.4 LLM Answer Generation',
        '   5.5 Conversation History & Context',
        '   5.6 Web Search Fallback',
        '6. Models & AI Components',
        '7. API Architecture (FastAPI)',
        '8. Metrics & Analytics System',
        '9. Response Caching',
        '10. Configuration & Environment',
        '11. Key Design Decisions & Trade-offs',
        '12. Data Flow Diagrams',
    ]
    for item in toc_items:
        add_para(doc, item)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. PROJECT OVERVIEW
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '1. Project Overview')
    add_para(doc,
        'The Ministry of Culture India Chatbot is a Retrieval-Augmented Generation (RAG) system '
        'that answers questions about Indian culture, heritage, museums, monuments, dance forms, '
        'schemes, and related topics. It uses data scraped from official government websites and '
        'answers queries in multiple languages (English, Hindi, Hinglish, and other Indian languages).'
    )
    add_para(doc, 'Core Technology Stack:', bold=True)

    tech_data = [
        ('Component', 'Technology', 'Purpose'),
        ('Web Framework', 'FastAPI (Python)', 'REST API + SSE streaming'),
        ('Vector Database', 'ChromaDB (PersistentClient)', 'Store and retrieve document embeddings'),
        ('Embedding Model', 'all-MiniLM-L6-v2 (sentence-transformers)', 'Generate 384-dim text embeddings'),
        ('LLM', 'qwen2.5:3b via Ollama (local)', 'Generate natural language answers'),
        ('Web Scraping', 'httpx + BeautifulSoup / Playwright', 'Crawl government websites'),
        ('Language Detection', 'langdetect + Hinglish word list', 'Identify query language'),
        ('Translation', 'deep_translator (Google Translate)', 'Multilingual support'),
        ('Web Search Fallback', 'Serper API', 'Real-time web search on trusted sites'),
        ('Metrics Storage', 'SQLite (WAL mode)', 'Analytics and usage tracking'),
    ]
    t = doc.add_table(rows=len(tech_data), cols=3)
    t.style = 'Table Grid'
    for i, row_data in enumerate(tech_data):
        for j, cell_text in enumerate(row_data):
            cell = t.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_bg(cell, '1a376c')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. HIGH-LEVEL ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '2. System Architecture — High-Level Flow')
    add_para(doc, 'The system has two independent phases: Ingestion (offline) and Query (online).')

    add_heading(doc, 'INGESTION PHASE (Offline / One-time)', level=2)
    add_code_block(doc, '''
+---------------------+       +-------------------+       +--------------------+
|   Government        |       |   Web Scraper     |       |   Text Processor   |
|   Websites          |  -->  |   (httpx /        |  -->  |   (chunk_size=800  |
|                     |       |    Playwright)    |       |    overlap=100)    |
| - indiaculture.gov  |       |                   |       |                    |
| - museumsofindia    |       | Extracts plain    |       | Splits into        |
| - asi.nic.in etc.   |       | text from HTML    |       | overlapping chunks |
+---------------------+       | or API JSON       |       +--------------------+
                              +-------------------+                |
                                                                   v
+---------------------+       +-------------------+       +--------------------+
|   ChromaDB          |       |   Embedding       |       |   Text Chunks      |
|   Vector Store      |  <--  |   Service         |  <--  |   (+ metadata:     |
|   (./data/chroma_db)|       |   (all-MiniLM-L6) |       |    url, title,     |
|                     |       |                   |       |    chunk_index)    |
|   893,000+ docs     |       |   384-dim vectors |       |                    |
+---------------------+       +-------------------+       +--------------------+
''')

    add_heading(doc, 'QUERY PHASE (Online / Real-time)', level=2)
    add_code_block(doc, '''
+------------------+     +------------------+     +------------------+
|   User Query     |     |  Language        |     |  Query           |
|   (any language) | --> |  Detection       | --> |  Embedding       |
|                  |     |  (langdetect +   |     |  (all-MiniLM)    |
+------------------+     |   Hinglish list) |     +------------------+
                         +------------------+              |
                                                           v
+------------------+     +------------------+     +------------------+
|   LLM Answer     |     |  Context         |     |  Vector Search   |
|   (qwen2.5:3b    | <-- |  Assembly        | <-- |  ChromaDB        |
|    via Ollama)   |     |  (top-k chunks   |     |  cosine sim      |
|                  |     |   + history)     |     |  threshold 0.55  |
+------------------+     +------------------+     +------------------+
         |
         v
+------------------+     +------------------+
|  If LLM says     |     |  Translation     |
|  "not found"     | --> |  (deep_translator|  --> Final Response
|  Web Fallback    |     |   if non-English)|
|  (Serper API)    |     +------------------+
+------------------+
''')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. DATA SOURCES
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '3. Data Sources')
    add_para(doc,
        'The knowledge base is built from official Ministry of Culture and ASI websites. '
        'Each source requires a different scraping strategy due to different rendering technologies.'
    )

    src_data = [
        ('Website', 'Rendering', 'Scraper', 'Content', 'Seed URLs'),
        ('indianculture.gov.in', 'React SPA', 'Playwright + API intercept', 'Dance, music, art, monuments, heritage', '25'),
        ('museumsofindia.gov.in', 'Server-rendered HTML', 'httpx + BeautifulSoup', 'Museum history, collections, timelines', '16 + auto-crawl'),
        ('indiaculture.gov.in', 'Server-rendered HTML', 'httpx + BeautifulSoup', 'Ministry schemes, news, tenders', 'Multiple'),
        ('asi.nic.in', 'Server-rendered HTML', 'httpx + BeautifulSoup', 'Archaeological sites, monuments', 'Multiple'),
    ]
    t = doc.add_table(rows=len(src_data), cols=5)
    t.style = 'Table Grid'
    for i, row_data in enumerate(src_data):
        for j, cell_text in enumerate(row_data):
            cell = t.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_bg(cell, '2e4a7a')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    add_heading(doc, 'Total Knowledge Base Size', level=2)
    add_para(doc, 'Collection: ministry_culture_kb')
    add_para(doc, 'Total documents (chunks): 893,037+')
    add_para(doc, 'Vector dimension: 384')
    add_para(doc, 'Storage: ./data/chroma_db (ChromaDB PersistentClient)')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. INGESTION PIPELINE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '4. Ingestion Pipeline')
    add_para(doc,
        'Ingestion is a 5-step offline process. The entry point is scripts/ingest_data.py '
        '(generic) or site-specific scripts (scrape_indianculture.py, scrape_museumsofindia.py). '
        'All scripts share the same service layer.'
    )

    add_code_block(doc, '''
Step 1: SCRAPE          Step 2: CHUNK          Step 3: EMBED
  |                       |                      |
  v                       v                      v
Fetch HTML/JSON  -->  Split into 800-char  -->  Generate 384-dim
from website          overlapping chunks        float32 vectors
                      (100 char overlap)        (all-MiniLM-L6-v2)
                             |
                             | attach metadata:
                             | {url, title, chunk_index,
                             |  chunk_length, source}
                             v
                      Step 4: STORE          Step 5: VERIFY
                        |                      |
                        v                      v
                    Add to ChromaDB  -->  Run test query
                    with upsert           check count
                    (MD5 chunk ID)        print stats
''')

    # 4.1 Web Scraping
    add_heading(doc, '4.1 Web Scraping', level=2)
    add_para(doc, 'Two scraping strategies are used depending on the site technology:', bold=False)

    add_heading(doc, 'Strategy A — httpx + BeautifulSoup (Server-rendered sites)', level=3)
    add_para(doc,
        'Used for sites like museumsofindia.gov.in where the server returns fully rendered HTML. '
        'Simpler, faster, and requires no browser automation.'
    )
    add_code_block(doc, '''
client = httpx.Client(verify=False, timeout=20, follow_redirects=True)
r = client.get(url)
soup = BeautifulSoup(r.text, "html.parser")

# Remove noise elements
for tag in ["script","style","nav","footer","header","noscript","iframe"]:
    soup(tag).decompose()

text = soup.get_text(separator=" ", strip=True)

# Auto-discover new links
for a in soup.find_all("a", href=True):
    full = urljoin(current_url, a["href"])
    if "museumsofindia.gov.in" in full and "/repository/" in full:
        to_visit.append(full)
''')

    add_heading(doc, 'Strategy B — Playwright + API Interception (React SPAs)', level=3)
    add_para(doc,
        'Used for indianculture.gov.in which is a React Single Page Application. '
        'Standard HTTP requests return empty HTML shells. Playwright runs a real Chromium browser '
        'and intercepts the JSON API calls the browser makes automatically to icvtestingold.nvli.in.'
    )
    add_code_block(doc, '''
async with async_playwright() as pw:
    browser = await pw.chromium.launch(headless=True)
    page    = await browser.new_page()

    # Intercept every HTTP response the page makes
    api_payloads = []
    async def on_response(response):
        if "icvtestingold.nvli.in" in response.url:
            if "json" in response.headers.get("content-type",""):
                data = await response.json()
                api_payloads.append(data)
    page.on("response", on_response)

    await page.goto(url, wait_until="domcontentloaded")
    await page.wait_for_load_state("networkidle", timeout=12000)
    await page.wait_for_timeout(2000)   # extra buffer for late API calls

    # Recursively extract text from captured JSON
    text = _extract_text_from_json(api_payloads)
''')

    add_para(doc,
        'The JSON extraction function recursively walks dict/list structures, '
        'pulling values from keys named: title, name, description, body, content, text, '
        'summary, detail, overview, caption, field_description, field_body, etc.'
    )

    # 4.2 Text Processing
    add_heading(doc, '4.2 Text Processing & Chunking', level=2)
    add_para(doc,
        'File: services/text_processor.py — TextProcessor class'
    )

    add_heading(doc, 'Text Cleaning Steps:', level=3)
    add_bullet(doc, 'HTML decoding: html.unescape() to convert &amp;, &nbsp; etc.')
    add_bullet(doc, 'HTML tag removal: BeautifulSoup with lxml parser')
    add_bullet(doc, 'Whitespace normalization: collapse multi-space/tabs/newlines')
    add_bullet(doc, 'Unicode cleanup: normalize unicode characters')
    add_bullet(doc, 'Navigation noise removal: strip boilerplate like "Skip to content", "Cookie Policy"')

    add_heading(doc, 'Chunking Algorithm:', level=3)
    add_para(doc, 'Parameters: chunk_size=800 chars, chunk_overlap=100 chars')
    add_code_block(doc, '''
Chunking with sentence-boundary alignment:

Raw text: "...sentence A. sentence B. sentence C. sentence D. sentence E..."
chunk_size = 800 characters, overlap = 100

Chunk 1: [  sentence A. sentence B. sentence C... ] <= 800 chars
         overlap:   [last 100 chars of chunk 1]
Chunk 2:            [ overlap + sentence D. sentence E... ] <= 800 chars

The algorithm finds the nearest sentence boundary (. ! ? |)
so chunks do not cut words in the middle.
Each chunk stores: text, url, title, start_pos, chunk_length
''')

    add_heading(doc, 'Output per page:', level=3)
    add_bullet(doc, 'A list of dicts: {text, url, title, start_pos, chunk_length}')
    add_bullet(doc, 'Average chunk size: ~400-600 chars after sentence-boundary alignment')
    add_bullet(doc, '~35 chunks per page on average (varies by content length)')

    # 4.3 Embedding
    add_heading(doc, '4.3 Embedding Generation', level=2)
    add_para(doc, 'File: services/embedding_service.py — EmbeddingService class (Singleton)')
    add_para(doc,
        'Model: all-MiniLM-L6-v2 from sentence-transformers library. '
        'This is a 22M parameter model optimized for semantic similarity tasks. '
        'It was trained on large NLI and STS datasets and produces embeddings '
        'in a 384-dimensional vector space.'
    )

    emb_data = [
        ('Property', 'Value'),
        ('Model name', 'all-MiniLM-L6-v2'),
        ('Parameters', '~22 million'),
        ('Output dimensions', '384'),
        ('Max input tokens', '256 (longer text is truncated)'),
        ('Inference', 'CPU (no GPU required)'),
        ('Library', 'sentence-transformers'),
        ('Normalization', 'L2 normalized (unit vectors)'),
    ]
    t = doc.add_table(rows=len(emb_data), cols=2)
    t.style = 'Table Grid'
    for i, row_data in enumerate(emb_data):
        for j, cell_text in enumerate(row_data):
            cell = t.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_bg(cell, '2e4a7a')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    add_code_block(doc, '''
# Batch embedding (used during ingestion)
texts      = [chunk["text"] for chunk in all_chunks]
embeddings = embedding_service.generate_embeddings_batch(
                 texts, batch_size=32, show_progress=True)
# Returns: List[List[float]]  shape: (N, 384)

# Single embedding (used at query time)
query_vector = embedding_service.generate_embedding(query_text)
# Returns: List[float]  shape: (384,)
''')

    add_para(doc,
        'The Singleton pattern ensures the model is loaded once per process '
        '(~2 seconds on first call) and reused for all subsequent requests.'
    )

    # 4.4 ChromaDB Storage
    add_heading(doc, '4.4 Vector Storage — ChromaDB', level=2)
    add_para(doc, 'File: services/vector_store.py — VectorStore class (Singleton)')
    add_para(doc,
        'ChromaDB is an open-source embedding database. The project uses PersistentClient '
        'which persists data to disk in SQLite + parquet format. All scrapers write to the '
        'same collection: ministry_culture_kb.'
    )

    add_code_block(doc, '''
Storage structure:
./data/chroma_db/
  |-- chroma.sqlite3       <- metadata, IDs, and document text
  |-- <uuid>/
       |-- data_level0.bin  <- HNSW index (for approximate nearest neighbour search)
       |-- header.bin
       |-- length.bin
       |-- link_lists.bin

Collection: ministry_culture_kb
Total documents: 893,037+

Document stored per chunk:
  ID       : MD5(url + start_pos)        <- deduplication key
  document : raw text of the chunk
  embedding: [384 float32 values]
  metadata : {
    url         : "https://museumsofindia.gov.in/repository/museum/...",
    title       : "Victoria Memorial Hall, Kolkata",
    chunk_index : 3,
    chunk_length: 756,
    source      : "museumsofindia.gov.in",
    ingested_at : "2024-01-15T14:32:11"
  }
''')

    add_heading(doc, 'Add operation (upsert):', level=3)
    add_code_block(doc, '''
# Batched in chunks of 5000 to avoid memory pressure
vector_store.add_documents(
    documents  = texts,       # List[str]
    embeddings = embeddings,  # List[List[float]]
    metadatas  = metadatas,   # List[Dict]
    ids        = ids          # List[str] — MD5 hash prevents duplicates
)
# ChromaDB does upsert: existing IDs are updated, new ones are inserted
''')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 5. RAG QUERY PIPELINE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '5. RAG Query-Time Pipeline')
    add_para(doc,
        'At query time, the system retrieves relevant chunks from ChromaDB using vector similarity, '
        'then passes those chunks as context to a local LLM to generate a grounded answer. '
        'This prevents hallucination — the LLM is explicitly told to answer only from the provided context.'
    )

    add_code_block(doc, '''
Complete query flow for /chat-hybrid-context endpoint:

User message
    |
    v
1. Language Detection (detect_language)
    |
    +-- English, Hindi, Hinglish, Tamil, Telugu, etc.
    |
    v
2. Query Translation to English (if not already English)
    |   deep_translator.GoogleTranslator().translate()
    |
    v
3. Query Embedding
    |   all-MiniLM-L6-v2.encode(english_query) -> 384-dim vector
    |
    v
4. Vector Search (ChromaDB)
    |   collection.query(query_embeddings=[vector], n_results=15)
    |   -> Returns 15 candidate chunks with distances
    |
    v
5. Similarity Filtering
    |   similarity = 1 / (1 + distance)
    |   Keep only chunks where similarity >= 0.55
    |
    v
6. LLM Context Assembly
    |   Top 5 chunks (similarity >= 0.65) passed to LLM
    |   Formatted as: "Source: <title>\n<snippet[:800]>"
    |
    v
7. LLM Answer Generation (qwen2.5:3b via Ollama)
    |   System prompt + conversation history + context + question
    |   num_predict=500, temperature=0.1, num_ctx=2048
    |
    v
8. Fallback Check
    |   If response contains "could not find verified information"
    |   -> Trigger web search fallback (Serper API)
    |
    v
9. Response Translation (if original query was non-English)
    |   deep_translator.GoogleTranslator().translate(answer, to=original_lang)
    |
    v
Final answer returned to user
''')

    # 5.1 API endpoints
    add_heading(doc, '5.1 Query Entry Points (API Endpoints)', level=2)

    ep_data = [
        ('Endpoint', 'Description', 'Vector?', 'History?', 'Streaming?'),
        ('POST /chat', 'Basic LLM chat (web search context)', 'No', 'No', 'No'),
        ('POST /chat/stream', 'Streaming version of /chat', 'No', 'No', 'Yes (SSE)'),
        ('POST /chat-context', 'Chat with conversation history', 'No', 'Yes', 'No'),
        ('POST /chat-hybrid', 'Vector search + web fallback', 'Yes', 'No', 'No'),
        ('POST /chat-hybrid-context', 'Vector + history (PRIMARY endpoint)', 'Yes', 'Yes', 'No'),
        ('POST /search', 'Semantic search (returns source list)', 'Yes', 'No', 'No'),
        ('GET /metrics/...', 'Usage analytics endpoints', '-', '-', '-'),
        ('GET /health', 'Health check', '-', '-', '-'),
    ]
    t = doc.add_table(rows=len(ep_data), cols=5)
    t.style = 'Table Grid'
    for i, row_data in enumerate(ep_data):
        for j, cell_text in enumerate(row_data):
            cell = t.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_bg(cell, '2e4a7a')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif i % 2 == 0:
                set_cell_bg(cell, 'F0F4FF')

    doc.add_paragraph()

    # 5.2 Language Detection
    add_heading(doc, '5.2 Language Detection & Translation', level=2)
    add_para(doc, 'File: services/language_service.py')
    add_code_block(doc, '''
Language Detection Flow:

Input text
    |
    v
Step 1: Check for Hinglish markers
    |   40+ Hindi words written in Latin script
    |   e.g. "kya", "hai", "mujhe", "batao", "museum", "kahan"
    |   If 2+ markers found -> return "hi" (Hindi/Hinglish)
    |
    v
Step 2: Try langdetect library
    |   langdetect.detect(text) -> ISO 639-1 language code
    |   e.g. "en", "hi", "ta", "te", "bn", "mr"
    |
    v
Step 3: Length check
    |   If text < 10 chars -> default to "en"
    |
    v
Language code returned (e.g. "en", "hi", "ta")

Translation (when language != "en"):
    Query:  deep_translator.GoogleTranslator(source="auto", target="en").translate(query)
    Answer: deep_translator.GoogleTranslator(source="en", target=lang).translate(answer)
''')

    # 5.3 Vector Search
    add_heading(doc, '5.3 Vector Search & Similarity Scoring', level=2)
    add_para(doc, 'File: services/vector_search_service.py — VectorSearchService class (Singleton)')
    add_code_block(doc, '''
Vector Search Process:

1. Query embedding: q_vec = model.encode(query)  # 384-dim vector

2. ChromaDB search:
   results = collection.query(
       query_embeddings = [q_vec],
       n_results        = 15,         # top-15 candidates
       include          = ["documents","metadatas","distances"]
   )

3. Similarity conversion:
   ChromaDB returns L2 distances (lower = more similar)
   We convert to similarity score:
       similarity = 1 / (1 + distance)
   This maps distance=0 -> similarity=1.0 (identical)
                   distance=1 -> similarity=0.5
                   distance=2 -> similarity=0.33

4. Filtering:
   Display threshold  : similarity >= 0.55  (shown in /search results)
   LLM context        : similarity >= 0.65  (passed to LLM for answers)

   Chunks below threshold are discarded — irrelevant context hurts LLM.

5. Return format:
   [
     {
       "text"      : "Victoria Memorial Hall is located in...",
       "title"     : "Victoria Memorial Hall",
       "url"       : "https://museumsofindia.gov.in/repository/museum/vmh_kol",
       "similarity": 0.78,
       "source"    : "museumsofindia.gov.in"
     },
     ...
   ]
''')

    # 5.4 LLM Answer Generation
    add_heading(doc, '5.4 LLM Answer Generation', level=2)
    add_para(doc, 'Files: services/vector_llm_service.py, services/llm_service.py, services/context_llm_service.py')
    add_para(doc,
        'The LLM used is qwen2.5:3b — a 3-billion parameter Qwen2.5 model served locally via Ollama. '
        'All inference happens on the same machine (CPU-only, no GPU required). '
        'The model is chosen for its balance of speed and accuracy at 3B parameter size.'
    )

    add_heading(doc, 'Ollama Configuration:', level=3)
    add_code_block(doc, '''
ollama.chat(
    model    = "qwen2.5:3b",
    messages = messages,
    options  = {
        "temperature": 0.1,    # Low temperature = factual, deterministic
        "num_predict": 500,    # Max output tokens (~375 words)
        "num_ctx"    : 2048,   # Context window (system + history + context + question)
        "num_thread" : 8,      # CPU threads
        "num_batch"  : 512,    # Processing batch size
        "num_gpu"    : 0       # Force CPU (no GPU)
    }
)
''')

    add_heading(doc, 'Message Structure sent to LLM:', level=3)
    add_code_block(doc, '''
messages = [
    {
        "role"   : "system",
        "content": SYSTEM_PROMPT    # 18 rules about domain, safety, language
    },
    {
        "role"   : "user",
        "content": "[prior turn 1]"
    },
    {
        "role"   : "assistant",
        "content": "[prior response 1]"
    },
    ...  (up to last 6 turns of history)
    {
        "role"   : "user",
        "content": "Question: What is the history of Salar Jung Museum?\n\n"
                   "Context:\n"
                   "Source: Salar Jung Museum\n"
                   "<chunk text 1[:800]>\n\n"
                   "Source: Museums of India\n"
                   "<chunk text 2[:800]>\n\n"
                   "...up to 5 sources...\n\n"
                   "Provide a detailed, comprehensive answer..."
    }
]
''')

    add_heading(doc, 'System Prompt Rules (summary):', level=3)
    rules = [
        '1. You are Ministry of Culture India AI assistant',
        '2. Answer ONLY from the provided context (no hallucination)',
        '3. If insufficient context, say "I could not find verified information"',
        '4. Never fabricate facts, dates, statistics, or names',
        '5. CRITICAL: Never use document metadata (Published Year, Size, SizeType) as actual facts',
        '6. Be detailed, factual, professional, and thorough',
        '7. Use the same language as the user',
        '8. For non-cultural/political queries, politely decline',
        '9. Format with bullet points for multi-part answers',
        '10. Always cite the source when providing specific facts',
    ]
    for rule in rules:
        add_bullet(doc, rule)

    # 5.5 Conversation History
    add_heading(doc, '5.5 Conversation History & Context', level=2)
    add_para(doc,
        'Context-aware endpoints (/chat-context, /chat-hybrid-context) maintain conversation history '
        'per session. History is stored client-side and sent with each request.'
    )
    add_code_block(doc, '''
Request body:
{
    "message"             : "Tell me more about its collection",
    "session_id"          : "abc123",
    "conversation_history": [
        {"role": "user",      "content": "What is Salar Jung Museum?"},
        {"role": "assistant", "content": "Salar Jung Museum is located in Hyderabad..."}
    ]
}

Server-side handling:
- Last 6 messages from history injected into LLM message list
- History-present queries skip the response cache (no stale responses)
- context_prompt_service.py builds history-aware user prompts
''')

    # 5.6 Web Search Fallback
    add_heading(doc, '5.6 Web Search Fallback', level=2)
    add_para(doc,
        'When the vector database cannot answer a question (no relevant chunks or LLM says '
        '"could not find verified information"), the system falls back to live web search '
        'using the Serper API, restricted to trusted Ministry of Culture websites.'
    )
    add_code_block(doc, '''
Web Search Fallback Chain:

Trigger conditions:
  (A) vector_results is empty (no chunks above threshold)
  (B) LLM response contains "could not find verified information"

Fallback steps:
  1. Call Serper API:
     POST https://google.serper.dev/search
     {
       "q"          : "<user question> site:indiaculture.gov.in OR site:asi.nic.in OR ...",
       "num"        : 5,
       "gl"         : "in",
       "hl"         : "en"
     }

  2. Extract organic results:
     [{title, link, snippet}, ...]

  3. Filter by trusted domains:
     TRUSTED_SITES = [
       "indiaculture.gov.in", "asi.nic.in", "indianculture.gov.in",
       "museumsofindia.gov.in", "ignca.gov.in", "ngmaindia.gov.in",
       "nationalmuseum.gov.in", "victoriamemorial-cal.org"
     ]

  4. Pass web results as context to LLM (same generation process)

  5. Translate response if needed

Last resort (web search also fails):
  "I could not find relevant information about this topic from trusted
   Ministry of Culture sources. Please try rephrasing your question."
''')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 6. MODELS & AI COMPONENTS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '6. Models & AI Components')

    add_heading(doc, '6.1 Embedding Model: all-MiniLM-L6-v2', level=2)
    add_para(doc,
        'This model converts text into dense 384-dimensional vectors. '
        'Similar texts produce vectors that are close together in vector space '
        '(measured by cosine similarity). It is used for BOTH ingestion (encoding '
        'document chunks) and querying (encoding the user question) — this symmetry '
        'is critical for accurate retrieval.'
    )
    add_bullet(doc, 'Architecture: 6-layer MiniLM transformer (distilled from larger models)')
    add_bullet(doc, 'Training: fine-tuned on Natural Language Inference (NLI) and Semantic Textual Similarity (STS) datasets')
    add_bullet(doc, 'Speed: ~50ms per query on CPU, ~2000 embeddings/second in batch mode')
    add_bullet(doc, 'Size: ~80MB on disk')

    add_heading(doc, '6.2 LLM: qwen2.5:3b via Ollama', level=2)
    add_para(doc,
        'Qwen2.5 is developed by Alibaba Cloud. The 3B variant is used for its balance '
        'of speed and quality. Ollama serves it locally as a REST API — no cloud API key needed.'
    )
    add_bullet(doc, 'Parameters: 3 billion')
    add_bullet(doc, 'Context window: 2048 tokens (configured), supports up to 32K natively')
    add_bullet(doc, 'Inference speed: ~15-25 tokens/second on CPU (response in 5-15 seconds)')
    add_bullet(doc, 'Temperature: 0.1 (near-deterministic for factual answers)')
    add_bullet(doc, 'Max output: 500 tokens (~375 words)')
    add_bullet(doc, 'Strengths: good English and multilingual understanding, follows instructions well')
    add_bullet(doc, 'Ollama endpoint: http://localhost:11434/api/chat')

    add_heading(doc, '6.3 Model Selection Logic at Runtime', level=2)
    add_code_block(doc, '''
At each request, the system checks which models are installed via:
  subprocess.run(["ollama", "list"], capture_output=True)

Priority order:
  1. qwen2.5:3b     - Primary (3B params, good accuracy)
  2. llama3.2:latest - Fallback (3B params, Meta)
  3. qwen2.5:3b     - Default if nothing found

Note: qwen2.5:1.5b and llama3.2:1b are commented out as
they were found too weak for complex information extraction.
''')

    add_heading(doc, '6.4 Translation: deep_translator (Google Translate)', level=2)
    add_para(doc,
        'deep_translator wraps Google Translate API for free-tier usage. '
        'It is used to translate both queries (to English before embedding) '
        'and final answers (from English back to the user\'s language).'
    )
    add_code_block(doc, '''
from deep_translator import GoogleTranslator

# Translate query to English
en_query = GoogleTranslator(source="auto", target="en").translate(hindi_query)

# Translate answer back
final = GoogleTranslator(source="en", target="hi").translate(en_answer)
''')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 7. API ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '7. API Architecture (FastAPI)')
    add_para(doc, 'File: app.py — FastAPI application entry point')
    add_code_block(doc, '''
FastAPI Application Structure:

app.py
  |-- CORS middleware (allow all origins for dev)
  |-- Routes registered:
  |     /api/v1/chat/*     <- routes/chat.py
  |     /api/v1/metrics/*  <- routes/metrics.py
  |
  +-- Startup event:
        - Loads VectorSearchService (loads ChromaDB + embedding model)
        - Loads MetricsService (opens SQLite)

routes/chat.py
  |-- POST /chat                   <- llm_service.generate_answer()
  |-- POST /chat/stream            <- SSE streaming with yield
  |-- POST /chat-context           <- context_llm_service.generate_context_aware_answer()
  |-- POST /chat-hybrid            <- vector_search + vector_llm_service
  |-- POST /chat-hybrid-context    <- vector_search + vector_llm_service + history
  |-- POST /search                 <- vector_search only (no LLM)

routes/metrics.py
  |-- GET  /metrics/summary        <- total requests, avg response time
  |-- GET  /metrics/top-topics     <- keyword-based topic classification
  |-- GET  /metrics/recent         <- latest N requests
  |-- GET  /metrics/by-language    <- breakdown by language
  |-- GET  /metrics/by-endpoint    <- breakdown by endpoint
''')

    add_heading(doc, 'Request/Response Models:', level=2)
    add_code_block(doc, '''
# Chat request
class ChatRequest(BaseModel):
    message             : str
    session_id          : str = ""
    conversation_history: List[Dict] = []

# Chat response
class ChatResponse(BaseModel):
    answer         : str
    language       : str
    response_time  : float
    sources        : List[Dict] = []

# Search request
class SearchRequest(BaseModel):
    query    : str
    max_results: int = 10

# Streaming (SSE) format
data: {"chunk": "Ministry of Culture", "done": false}
data: {"chunk": " India was...", "done": false}
data: {"done": true, "full_response": "..."}
''')

    add_heading(doc, 'CORS & Security:', level=2)
    add_bullet(doc, 'CORS: Allow all origins (development setting — restrict in production)')
    add_bullet(doc, 'No authentication on API (intended for internal use)')
    add_bullet(doc, 'Input validation via Pydantic models')
    add_bullet(doc, 'Session IDs used for metrics tracking (no server-side session state)')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 8. METRICS & ANALYTICS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '8. Metrics & Analytics System')
    add_para(doc, 'File: services/metrics_service.py — MetricsService class (Singleton)')
    add_para(doc,
        'Every API request is logged to a SQLite database with WAL (Write-Ahead Logging) '
        'mode for concurrent read/write without locking.'
    )

    add_heading(doc, 'Database Schema:', level=2)
    add_code_block(doc, '''
Table: requests
  id            INTEGER PRIMARY KEY AUTOINCREMENT
  timestamp     TEXT    NOT NULL      -- ISO 8601 datetime
  endpoint      TEXT    NOT NULL      -- e.g. "chat-hybrid-context"
  response_time REAL                  -- seconds (float)
  session_id    TEXT
  client_ip     TEXT
  language      TEXT                  -- detected language code
  query_text    TEXT                  -- English version of query
  feedback      INTEGER               -- 1=positive, -1=negative, NULL=not rated
''')

    add_heading(doc, 'Topic Classification:', level=2)
    add_para(doc,
        '15 topic categories classified by keyword matching on the stored English query text. '
        'Each query is matched against keyword lists; the first matching category wins.'
    )
    categories = [
        'Historical Monuments (monument, fort, temple, mahal, minar, qutub, taj, hampi...)',
        'Museums & Galleries (museum, gallery, exhibit, artifact, salar jung, ngma...)',
        'Classical Dance Forms (dance, bharatnatyam, kathak, odissi, kuchipudi...)',
        'Vedic Heritage (vedic, veda, upanishad, sanskrit, rigveda...)',
        'UNESCO World Heritage (unesco, world heritage, intangible heritage...)',
        'National Archives (archive, manuscript, document, record, abhilekh...)',
        'Freedom Movement (freedom, independence, gandhi, nehru, 1857...)',
        'Art & Craft (art, craft, painting, sculpture, warli, madhubani...)',
        'Festivals & Events (festival, mela, utsav, diwali, holi, navratri...)',
        'Music (raga, classical music, carnatic, hindustani, tabla, sitar...)',
        'Schemes & Grants (scheme, grant, scholarship, fellowship, award...)',
        'Tenders & Procurement (tender, bid, procurement, rfp, quotation...)',
        'Personalities (biography, artist, author, poet, sculptor...)',
        'Architecture (architectural style, structure, indo-saracenic, mughal...)',
        'Literature (poetry, novel, book, author, literary, language...)',
    ]
    for cat in categories:
        add_bullet(doc, cat)

    add_heading(doc, 'Analytics Endpoints:', level=2)
    add_code_block(doc, '''
GET /metrics/summary
  -> { total_requests, avg_response_time, requests_today, unique_sessions }

GET /metrics/top-topics?limit=10
  -> [ { topic: "Museums & Galleries", count: 142 }, ... ]

GET /metrics/recent?limit=20
  -> [ { timestamp, endpoint, language, query_text, response_time }, ... ]

GET /metrics/by-language
  -> [ { language: "en", count: 890 }, { language: "hi", count: 234 }, ... ]

GET /metrics/by-endpoint
  -> [ { endpoint: "chat-hybrid-context", count: 567 }, ... ]
''')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 9. RESPONSE CACHING
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '9. Response Caching')
    add_para(doc,
        'To avoid redundant LLM calls for identical queries, the system uses in-memory '
        'response caches. Each cache is an ordered dict limited to 100 entries (LRU eviction).'
    )
    add_code_block(doc, '''
Cache key generation:
  combined = f"{endpoint}:{question.lower().strip()}:{language}:{history_length}"
  key      = MD5(combined)

Cache behavior:
  - Cache HIT  : return cached answer immediately (no LLM call)
  - Cache MISS : generate answer, store in cache, return

Cache invalidation:
  - Size limit: 100 entries per cache (oldest entry evicted when full)
  - Session-aware: queries with conversation_history (history_len > 0) BYPASS cache
    (personalized responses should not be cached across sessions)
  - Process restart: cache cleared (in-memory only, not persisted)

Separate caches per service:
  vector_llm_service.py  : _vector_response_cache  (100 entries)
  llm_service.py         : _response_cache          (100 entries)
  context_llm_service.py : _context_response_cache  (100 entries)
''')

    # ══════════════════════════════════════════════════════════════════════════
    # 10. CONFIGURATION & ENVIRONMENT
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '10. Configuration & Environment')

    env_data = [
        ('Variable', 'Default', 'Purpose'),
        ('TEST_MODE', 'unset', 'If set to "1", uses ./data/chroma_db_test instead of prod DB'),
        ('SERPER_API_KEY', 'required', 'API key for Serper (Google Search) web fallback'),
        ('OLLAMA_HOST', 'localhost:11434', 'Ollama server address'),
    ]
    t = doc.add_table(rows=len(env_data), cols=3)
    t.style = 'Table Grid'
    for i, row_data in enumerate(env_data):
        for j, cell_text in enumerate(row_data):
            cell = t.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_bg(cell, '2e4a7a')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    add_heading(doc, 'Startup Commands:', level=2)
    add_code_block(doc, '''
# Install dependencies
pip install -r requirements.txt
pip install playwright && playwright install chromium

# Install Ollama and model
# Download from ollama.com, then:
ollama pull qwen2.5:3b

# Run ingestion (one-time)
python scripts/scrape_indianculture.py
python scripts/scrape_museumsofindia.py
python scripts/ingest_data.py          # other sites

# Start API server
uvicorn app:app --host 0.0.0.0 --port 8000 --reload

# Dev mode with auto-reload
python -m uvicorn app:app --reload
''')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 11. DESIGN DECISIONS & TRADE-OFFS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '11. Key Design Decisions & Trade-offs')

    decisions = [
        (
            'Local LLM (Ollama) instead of OpenAI/Claude API',
            'No API costs, data stays on-premise, no internet dependency for LLM',
            'Slower inference (CPU-only), smaller model capability'
        ),
        (
            'ChromaDB over Pinecone/Weaviate',
            'Free, open-source, runs locally, no cloud dependency',
            'Single-node only, no distributed scaling'
        ),
        (
            'all-MiniLM-L6-v2 over larger embedding models',
            'Fast (80MB), good quality for English/multilingual, widely tested',
            'Max 256 tokens input — very long chunks get truncated'
        ),
        (
            'Similarity threshold 0.55 (display) / 0.65 (LLM)',
            'Filters irrelevant results before they confuse the LLM',
            'May miss some relevant content if phrased very differently from stored chunks'
        ),
        (
            'Web search fallback (Serper)',
            'Handles topics not in the knowledge base (current events, new tenders)',
            'Trusted-site filtering may miss some official content'
        ),
        (
            'Chunk size 800 chars with 100 overlap',
            'Balances context richness vs. embedding quality (MiniLM max 256 tokens)',
            'Long multi-part answers may be split across chunks'
        ),
        (
            'Playwright for React SPAs',
            'Correctly captures JavaScript-rendered content',
            'Slower than httpx (2-5s per page), requires Chromium install'
        ),
        (
            'MD5 chunk ID for deduplication',
            'Re-running scrapers is safe — existing chunks are updated, not duplicated',
            'MD5 collision is theoretically possible (negligible in practice)'
        ),
    ]

    for decision, pro, con in decisions:
        add_para(doc, decision, bold=True)
        p = doc.add_paragraph()
        p.add_run('Rationale: ').bold = True
        p.add_run(pro)
        p2 = doc.add_paragraph()
        p2.add_run('Trade-off: ').bold = True
        p2.add_run(con)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 12. FULL FLOW DIAGRAMS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '12. Full Flow Diagrams')

    add_heading(doc, '12.1 Complete Ingestion Pipeline', level=2)
    add_code_block(doc, '''
                    INGESTION PIPELINE
                    ==================

 [Government Websites]
         |
         |  HTTP GET or Playwright headless browser
         |
         v
 +--------------------+
 |   Raw HTML / JSON  |
 +--------------------+
         |
         |  BeautifulSoup / JSON parser
         |  - Remove: <script>, <style>, <nav>, <footer>
         |  - Extract: visible text / JSON text fields
         v
 +--------------------+
 |   Clean Plain Text |
 |   (per page)       |
 +--------------------+
         |
         |  TextProcessor.chunk_text()
         |  - Split by 800 chars
         |  - Align to sentence boundaries
         |  - 100-char overlap between chunks
         v
 +--------------------+     metadata: {
 |   Text Chunks      |       url, title, source,
 |   (N per page)     |       chunk_index, chunk_length
 +--------------------+     }
         |
         |  EmbeddingService.generate_embeddings_batch()
         |  - Model: all-MiniLM-L6-v2
         |  - Batch size: 32
         |  - Output: 384-dim float32 vectors
         v
 +--------------------+
 |   Float Vectors    |
 |   shape: (N, 384)  |
 +--------------------+
         |
         |  VectorStore.add_documents()
         |  - ID: MD5(url + start_pos)
         |  - Batch insert: 5000 per batch
         |  - Upsert semantics (deduplication)
         v
 +---------------------------+
 |   ChromaDB                |
 |   Collection:             |
 |   ministry_culture_kb     |
 |   893,037+ documents      |
 |   ./data/chroma_db        |
 +---------------------------+
''')

    add_heading(doc, '12.2 Complete Query Pipeline (chat-hybrid-context)', level=2)
    add_code_block(doc, '''
                    QUERY PIPELINE (/chat-hybrid-context)
                    ======================================

 [User]  "Tell me about Ellora Caves" (in Hindi: "एलोरा गुफाओं के बारे में बताएं")
         |
         v
 [FastAPI /chat-hybrid-context]
         |
         |  detect_language(text)
         |  -> "hi"  (Hindi detected)
         v
 [GoogleTranslator]
         |  translate("hi" -> "en")
         |  -> "Tell me about Ellora Caves"
         v
 [EmbeddingService]
         |  all-MiniLM-L6-v2.encode("Tell me about Ellora Caves")
         |  -> [0.023, -0.145, 0.089, ...]  (384 values)
         v
 [ChromaDB collection.query()]
         |  n_results = 15
         |  -> Returns 15 candidate chunks with L2 distances
         v
 [Similarity Filtering]
         |  similarity = 1 / (1 + distance)
         |  Display: keep >= 0.55
         |  LLM input: keep >= 0.65  (top quality only)
         v
 [Context Assembly]
         |  Top 5 chunks formatted as:
         |  "Source: Ajanta and Ellora Caves\n<text[:800]>"
         |  + Last 6 messages from conversation_history
         v
 [Ollama qwen2.5:3b]
         |  System: SYSTEM_PROMPT (18 rules)
         |  History: last 6 messages
         |  User: question + context
         |  Options: temp=0.1, num_predict=500
         |
         |  -> English answer generated
         v
 [Fallback Check]
         |  answer.lower() contains "could not find verified information"?
         |  YES -> Serper API web search -> new LLM call
         |  NO  -> proceed
         v
 [GoogleTranslator]
         |  translate(english_answer, "en" -> "hi")
         |  -> Hindi answer
         v
 [MetricsService.record_request()]
         |  Store to SQLite: endpoint, time, language, query
         v
 [Response to User]
         {
           "answer"       : "एलोरा गुफाएं...",
           "language"     : "hi",
           "response_time": 4.2,
           "sources"      : [{title, url, similarity}, ...]
         }
''')

    add_heading(doc, '12.3 Web Search Fallback Flow', level=2)
    add_code_block(doc, '''
                    WEB SEARCH FALLBACK
                    ===================

         Trigger: No vector results OR LLM says "could not find"
                              |
                              v
                    [Serper API Request]
                    POST google.serper.dev/search
                    {
                      "q"  : "<question> site:indiaculture.gov.in
                               OR site:asi.nic.in OR site:museumsofindia.gov.in
                               OR site:ignca.gov.in OR site:nationalmuseum.gov.in",
                      "num": 5,
                      "gl" : "in"
                    }
                              |
                              v
                    [Parse organic results]
                    [{title, link, snippet}, ...]
                              |
                              v
                    [Filter by TRUSTED_SITES list]
                    Keep only official government/cultural domains
                              |
                              v
                    [LLM generation with web context]
                    Same as normal flow but context = web snippets
                              |
                              v
                    [Translate if needed + Return]
''')

    add_heading(doc, '12.4 System Component Map', level=2)
    add_code_block(doc, '''
                    SYSTEM COMPONENT MAP
                    ====================

  +-----------+    +-----------+    +-----------+    +-----------+
  |  Frontend |    | FastAPI   |    | Services  |    | External  |
  |  (React/  |    | app.py    |    | Layer     |    | Systems   |
  |   Flutter)|    |           |    |           |    |           |
  +-----------+    +-----------+    +-----------+    +-----------+
        |               |                |                |
        | HTTP POST      |                |                |
        |-------------->|                |                |
        |               |                |                |
        |               |-- language  -->|                |
        |               |   _service     |                |
        |               |                |                |
        |               |-- embedding -->|                |
        |               |   _service     |                |
        |               |                |                |
        |               |-- vector    -->|                |
        |               |   _search      |-- ChromaDB    |
        |               |   _service     |   ./data/     |
        |               |                |   chroma_db   |
        |               |-- vector    -->|               |
        |               |   _llm         |-- Ollama:11434|
        |               |   _service     |   qwen2.5:3b  |
        |               |                |                |
        |               |-- (fallback) ->|-- Serper API  |
        |               |   web_search   |   (internet)  |
        |               |                |                |
        |               |-- metrics   -->|-- SQLite      |
        |               |   _service     |   metrics.db  |
        |               |                |                |
        |<-- response --|                |                |
''')

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # CLOSING
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, 'Interview Quick-Reference', level=1)
    add_para(doc, 'Likely interview questions and concise answers:', bold=True)

    qa = [
        (
            'Q: What is RAG and how does this project implement it?',
            'A: RAG = Retrieval-Augmented Generation. Instead of relying on the LLM\'s training data, '
            'we retrieve relevant document chunks from ChromaDB using vector similarity search, '
            'then inject those chunks as context into the LLM prompt. The LLM is instructed to answer '
            'ONLY from the provided context, preventing hallucination.'
        ),
        (
            'Q: How does vector similarity search work?',
            'A: Text is converted to 384-dim float vectors using all-MiniLM-L6-v2. '
            'ChromaDB stores all chunk vectors. At query time, the question is also embedded '
            'and ChromaDB finds the nearest neighbours using L2 distance. '
            'We convert L2 distance to similarity: similarity = 1/(1+distance). '
            'Chunks with similarity >= 0.65 are passed to the LLM.'
        ),
        (
            'Q: Why local LLM instead of GPT-4?',
            'A: Cost (no per-token fees), privacy (data stays on-premise), and '
            'no internet dependency for inference. qwen2.5:3b is adequate for '
            'factual Q&A when good context is provided.'
        ),
        (
            'Q: How do you handle non-English queries?',
            'A: Three-step: (1) detect language using langdetect + custom Hinglish word list, '
            '(2) translate query to English for vector search (all-MiniLM-L6-v2 works best in English), '
            '(3) translate the English answer back to the original language using Google Translate.'
        ),
        (
            'Q: What prevents the LLM from hallucinating?',
            'A: The system prompt contains explicit rules: "Answer ONLY from provided context", '
            '"Never fabricate facts", "CRITICAL: Never use document metadata as facts". '
            'Temperature is set to 0.1 (near-deterministic). '
            'When context is insufficient, the LLM is instructed to say "I could not find verified information" '
            'which triggers the web search fallback.'
        ),
        (
            'Q: How does the web search fallback work?',
            'A: When vector results are absent or insufficient, Serper API searches Google '
            'restricted to trusted official domains (indiaculture.gov.in, asi.nic.in, etc.). '
            'The search snippets are then used as context for another LLM generation call.'
        ),
        (
            'Q: How is the knowledge base kept up to date?',
            'A: Re-running the scraper scripts updates the ChromaDB collection. '
            'MD5 chunk IDs ensure deduplication — existing chunks are updated (upsert), '
            'new chunks are added. No manual deletion needed.'
        ),
    ]

    for q, a in qa:
        add_para(doc, q, bold=True, color=(0x1a, 0x37, 0x6c))
        add_para(doc, a)
        doc.add_paragraph()

    # Save
    output_path = './Ministry_Culture_Chatbot_Technical_Architecture.docx'
    doc.save(output_path)
    print(f"[OK] Document saved: {output_path}")


if __name__ == "__main__":
    build_doc()
